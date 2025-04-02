"""
Модуль для отримання повідомлень з заданих телеграм-каналів за вказаний діапазон дат,
фільтрації повідомлень за ключовими словами (regex) та генерації звіту у форматі DOCX
з використанням бібліотеки python-docx.

Приклад файлу config.json:
{
  "api_id": 123456,
  "api_hash": "your_api_hash_here",
  "channels": [
    {"Назва групи каналів-1": ["channel1", "channel2"]},
    {"Назва групи каналів-2": ["channel3", "channel4"]}
  ],
  "start_date": "02.04.2025",
  "end_date": "02.04.2025",
  "keywords": ["regex1", "regex2"]
}
"""

import asyncio
import json
import logging
import re
from datetime import datetime

from telethon import TelegramClient

from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Налаштування логування
logging.basicConfig(level=logging.INFO, format='[%(levelname)s] %(asctime)s - %(message)s')


class TelegramSpider:
    def __init__(self, config_path: str = 'config.json'):
        """
        Ініціалізація спайдера із завантаженням конфігурації.
        """
        self.logger = logging.getLogger(self.__class__.__name__)
        self.config_path = config_path
        self.config = self._load_config()
        self.client = TelegramClient(
            'session_name',
            self.config.get('api_id'),
            self.config.get('api_hash')
        )
        # Структура для збереження повідомлень за групами каналів
        self.results_by_group = {}

    def _load_config(self) -> dict:
        """
        Завантаження конфігурації з файлу config.json.
        """
        try:
            with open(self.config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
            self.logger.info("Конфігурацію завантажено успішно.")
            return config
        except Exception as e:
            self.logger.error(f"Помилка завантаження конфігурації: {e}")
            raise

    async def authenticate(self):
        """
        Підключення та авторизація в Telegram.
        """
        self.logger.info("Підключення до Telegram...")
        await self.client.start()
        self.logger.info("Авторизація пройшла успішно.")

    async def fetch_messages(self):
        """
        Отримання повідомлень з каналів за заданих груп та фільтрація їх за діапазоном дат і ключовими словами.
        """
        try:
            start_date = datetime.strptime(self.config.get('start_date'), '%d.%m.%Y')
            end_date = datetime.strptime(self.config.get('end_date'), '%d.%m.%Y')
        except Exception as e:
            self.logger.error(f"Невірний формат дат у конфігурації: {e}")
            raise
        keywords = self.config.get('keywords', [])
        groups = self.config.get('channels', [])
        # Перебір груп каналів
        for group in groups:
            for group_name, channel_list in group.items():
                self.logger.info(f"Обробка групи: {group_name}")
                if group_name not in self.results_by_group:
                    self.results_by_group[group_name] = []
                # Перебір каналів у групі
                for channel in channel_list:
                    self.logger.info(f"Отримання повідомлень з каналу: {channel}")
                    try:
                        async for message in self.client.iter_messages(channel, offset_date=end_date, reverse=True):
                            message_date = message.date.replace(tzinfo=None)
                            if message_date < start_date:
                                break
                            if message.message:
                                for pattern in keywords:
                                    if re.search(pattern, message.message, re.IGNORECASE):
                                        lines = message.message.splitlines()
                                        title = lines[0] if lines else ""
                                        content = "\n".join(lines[1:]) if len(lines) > 1 else ""
                                        self.results_by_group[group_name].append({
                                            'channel': channel,
                                            'id': message.id,
                                            'title': title,
                                            'content': content,
                                            'date': message_date.strftime('%d.%m.%Y %H:%M')
                                        })
                                        break
                    except Exception as e:
                        self.logger.error(f"Помилка при отриманні повідомлень з каналу {channel}: {e}")

    def _add_hyperlink(self, paragraph, url, text):
        """
        Додає гіперпосилання до абзацу з синім підкресленим текстом.
        """
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        # Встановлення підкреслення
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        rPr.append(underline)
        # Встановлення кольору тексту (синій - 0000FF)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        new_run.append(rPr)
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink

    def generate_html_report(self, output_file: str = None):
        """
        Генерація звіту у форматі DOCX з використанням python-docx.
        """
        self.logger.info("Генерація звіту у форматі DOCX...")
        document = Document()
        # Налаштування полів документа
        for section in document.sections:
            section.top_margin = Mm(10)
            section.bottom_margin = Mm(10)
            section.left_margin = Mm(25)
            section.right_margin = Mm(15)
        # Налаштування шрифту документа
        style = document.styles['Normal']
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(14)
        # Перебір груп повідомлень
        for group, messages in self.results_by_group.items():
            # Заголовок групи: жирним, по центру
            p_group = document.add_paragraph()
            p_group.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_group = p_group.add_run(group)
            run_group.bold = True
            run_group.font.name = "Times New Roman"
            run_group.font.size = Pt(14)
            # Перебір повідомлень у групі
            for msg in messages:
                # Рядок з посиланням на джерело
                p_link = document.add_paragraph()
                source_url = f"https://t.me/{msg['channel']}/{msg['id']}"
                self._add_hyperlink(p_link, source_url, source_url)
                run_info = p_link.add_run(msg['title'])  #  {msg['date']}
                run_info.font.name = "Times New Roman"
                run_info.font.size = Pt(14)
                # Рядок із текстом статті з вирівнюванням по ширині
                p_content = document.add_paragraph()
                p_content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run_content = p_content.add_run(msg['content'])
                run_content.font.name = "Times New Roman"
                run_content.font.size = Pt(14)
                # додаємо порожній рядок для розділення записів
                document.add_paragraph("")
        # Формування імені вихідного файлу за замовчуванням, якщо не задано output_file
        start_date = self.config.get('start_date')
        end_date = self.config.get('end_date')
        if output_file is None:
            if start_date == end_date:
                output_file = f"звіт_{start_date}.docx"
            else:
                output_file = f"звіт_{start_date}-{end_date}.docx"
        try:
            document.save(output_file)
            self.logger.info(f"Звіт збережено у файлі: {output_file}")
        except Exception as e:
            self.logger.error(f"Помилка при збереженні звіту: {e}")

    async def run(self):
        """
        Підключення та отримання звіту
        """
        await self.authenticate()  # Авторизація
        await self.fetch_messages()  # Отримання та фільтрація повідомлень
        self.generate_html_report()  # Генерація звіту у форматі DOCX


if __name__ == "__main__":
    spider = TelegramSpider('config.json')
    loop = asyncio.get_event_loop()
    loop.run_until_complete(spider.run())
